'''=========================THE OBJECTIVE======================================
Assume you are the corporate mole in the episode and use Python to hide a 
secret message summarizing bid details within an official-looking text document. 
Start with an unencrypted message and finish with an encrypted version.'''


import docx
from docx.shared import RGBColor, Pt


def secret_message(file):
    '''Loads a file that contains secret message and returns list of paragraphs'''
    secret = docx.Document(file)
    secret_list = []
    for paragraph in secret.paragraphs:
        secret_list.append(paragraph.text)

    return secret_list


def real_message(file):
    '''Loads a file that contains real message and returns a list of paragraphs'''
    real = docx.Document(file)
    real_list = []
    for paragraph in real.paragraphs:
        if len(paragraph.text) > 0:
            real_list.append(paragraph.text)
        
    return real_list


def template(file):
    '''Loads a template file for final message along with some text'''
    doc = docx.Document(file)
    doc.add_heading('Morland Holmes', 0)
    subtitle = doc.add_heading('Global Consulting & Negotiations')
    subtitle.alignment = 1
    doc.add_paragraph('December 17, 2015')
    doc.add_paragraph('')

    return doc


def set_spacing(paragraph):
    """Use docx to set line spacing between paragraphs."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def main():
    secret = secret_message(file_2)
    real = real_message(file_3)
    doc = template(file_1)
    length_real = len(real)
    count_real = 0 # index of current line in real (hidden) message

    for line in secret:
        if count_real < length_real and line == '':
            paragraph = doc.add_paragraph(real[count_real])
            paragraph_index = len(doc.paragraphs) - 1
            run = doc.paragraphs[paragraph_index].runs[0]
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255) # make it red to test
            count_real += 1
        else:
             paragraph = doc.add_paragraph(line)
            
    set_spacing(paragraph)

    doc.save('chapter_6\\ciphertext_message_letterhead.docx')
    


if __name__ == "__main__":
    print('starting encoder...')
    file_1 = 'chapter_6\\template.docx'
    file_2 = 'chapter_6\\fakeMessage.docx'
    file_3 = 'chapter_6\\realMessage.docx'

    main()
    print('Finsihed')

